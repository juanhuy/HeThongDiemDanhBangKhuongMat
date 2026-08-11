"""Dịch vụ AI cho hệ thống tài liệu (chạy offline, không cần API ngoài).

- Trích xuất văn bản từ PDF (.pdf) / Word (.docx).
- Sinh thẻ (tags), từ khóa (keywords), tóm tắt (summary), ý chính (key_points).
- Sinh flashcard dạng câu hỏi - trả lời từ nội dung.
"""
import re
import unicodedata

# Vietnamese stopwords (tập con thông dụng)
VIETNAMESE_STOPWORDS = set("""
và là của có trong cho các với như không được một những khi này đó theo tại thì ra
sẽ về nếu còn rồi từ đó để do cũng nhưng hoặc hay nên vì mà bởi cả đã đang sẽ là
về phía trên dưới sau trước giữa cùng về lại lên xuống vào ra qua khỏi theo tại
đối với bởi vì ngoài ra hơn nhất mỗi mọi vài những các đây đó kia này ấy có thể
được phải cần nên sẽ nếu thì hay hoặc chúng tôi chúng ta anh chi em bạn thầy cô
sinh viên giảng viên học phần môn học điểm danh buổi học lớp học hệ thống phần mềm
bài giảng tài liệu nội dung chương trình đề tài nghiên cứu khoa học công nghệ thông tin
nhau cũng đều làm cho bi de toi ta minh noi lai roi thi thi va vd vv diem
also the and of to in for on with as by at from this that these those is are was were be been
will would can could should must have has had do does did not no yes your our their its
""".split())


def _normalize_keyword(word: str) -> str:
    """Chuẩn hóa: bỏ dấu để gom từ cùng gốc, hạ thường, bỏ ký tự đặc biệt."""
    word = unicodedata.normalize("NFD", word.lower().strip())
    word = "".join(c for c in word if unicodedata.category(c) != "Mn")
    return word


def extract_text(data: bytes, ext: str) -> str:
    """Trích xuất văn bản thuần từ nội dung file."""
    ext = (ext or "").lower().strip(".")
    try:
        if ext == "pdf":
            return _extract_pdf(data)
        if ext in ("docx", "doc"):
            return _extract_docx(data)
    except Exception:
        pass
    return ""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _split_sentences(text: str) -> list:
    """Tách câu tiếng Việt (theo dấu chấm / chấm than / chấm hỏi)."""
    raw = re.split(r"(?<=[.!?])\s+|\n+", text)
    sentences = []
    for s in raw:
        s = s.strip()
        if len(s) >= 15:
            sentences.append(s)
    return sentences


def _word_frequencies(text: str, limit: int = 20) -> list:
    """Đếm tần suất từ (bỏ stopwords + từ quá ngắn/dài). Trả về [{word, key, count}]."""
    words = re.findall(r"[A-Za-zÀ-ỹ0-9]+", text.lower())
    freq = {}          # normalized_key -> {"count": int, "display": str}
    for w in words:
        if len(w) < 4 or len(w) > 40:
            continue
        key = _normalize_keyword(w)
        if key in VIETNAMESE_STOPWORDS or key == "ve":
            continue
        entry = freq.setdefault(key, {"count": 0, "display": w})
        entry["count"] += 1
        # Giữ dạng hiển thị dài hơn (nhiều dấu) làm đại diện cho nhóm từ cùng gốc
        if len(w) > len(entry["display"]):
            entry["display"] = w
    return sorted(
        [{"word": v["display"], "key": k, "count": v["count"]} for k, v in freq.items()],
        key=lambda x: -x["count"],
    )[:limit]


def _restore_keyword(key: str) -> str:
    return key.upper()


def analyze_document(text: str, max_text: int = 45000):
    """Trả về dict: tags, keywords, summary, key_points, chapters từ văn bản trích xuất."""
    raw = (text or "")[:max_text]
    text = _clean_text(raw)
    if not text:
        return {
            "tags": "", "keywords": "", "summary": "",
            "key_points": "", "chapters": [], "content_text": raw,
        }

    top_words = _word_frequencies(text, limit=12)
    tags = [w["word"] for w in top_words[:5]]
    keywords = [w["word"] for w in top_words[:8]]

    sentences = _split_sentences(text)
    summary = ""
    key_points = []

    if sentences:
        # TÓM TẮT TRÍCH RÚT (extractive): chấm điểm câu theo từ khóa + vị trí
        def _score_sentence(s, idx):
            sl = _normalize_keyword(s)
            score = sum(1 for w in top_words[:8] if w["key"] in sl)
            if idx == 0:
                score += 2  # câu mở đầu thường là giới thiệu
            if len(s) < 40:
                score -= 1
            if len(s) > 300:
                score -= 1
            return score

        ranked = sorted(((i, _score_sentence(s, i)) for i, s in enumerate(sentences)), key=lambda x: -x[1])
        # Chọn 3-4 câu điểm cao, giữ thứ tự xuất hiện
        best_idx = sorted(i for i, _ in ranked[:3])
        summary = " ".join(sentences[i] for i in best_idx)

        # Ý chính: chọn câu chứa nhiều từ khóa nhất
        scored = []
        for i, s in enumerate(sentences):
            sl = _normalize_keyword(s)
            score = sum(1 for w in top_words[:8] if w["key"] in sl)
            scored.append((score, s))
        scored.sort(key=lambda x: -x[0])
        seen = set()
        for score, s in scored:
            if score < 1:
                break
            if s in seen:
                continue
            seen.add(s)
            key_points.append(s)
            if len(key_points) >= 5:
                break
        if not key_points and sentences:
            key_points = sentences[:3]

    # Tóm tắt cấu trúc theo chương/mục (tách từ văn bản GỐC để giữ dòng tiêu đề)
    chapters = []
    for ch in extract_chapters(raw):
        body = _clean_text(" ".join(ch["body"]))
        if len(body.strip()) < 30:
            continue
        kw = _word_frequencies(body, limit=5)
        ch_sentences = _split_sentences(body)
        chapters.append({
            "title": ch["title"],
            "summary": ch_sentences[0] if ch_sentences else body[:150],
            "keywords": [w["word"] for w in kw],
        })

    # Thuật ngữ (fallback): từ các câu định nghĩa kiểu 'X là ...'
    terms = []
    seen_terms = set()
    for s in sentences:
        dc = _definition_card(s)
        if dc:
            term = dc["question"].replace(" là gì?", "").strip()
            if term and term not in seen_terms and len(term) <= 80:
                seen_terms.add(term)
                terms.append({"term": term, "definition": s})

    return {
        "tags": ", ".join(tags),
        "keywords": ", ".join(keywords),
        "summary": summary,
        "key_points": "\n".join(key_points),
        "chapters": chapters,
        "conclusion": sentences[-1] if sentences else "",
        "terms": terms,
        "content_text": raw,
    }


# =========================================================================
# TRÍCH XUẤT CẤU TRÚC CHƯƠNG / MỤC
# =========================================================================
_HEADING_PATTERNS = [
    re.compile(r"^(chương|chuong|phần|phan|mục|muc|bài|bai)\s+([0-9]+|[ivxlcdm]+)\b", re.IGNORECASE),
    re.compile(r"^([0-9]+(\.[0-9]+){0,2})[\s．\.\)\-–—]\s+[A-ZÀ-ỸẠ-Ỵ]"),
    re.compile(r"^(mở đầu|giới thiệu|tổng quan|kết luận|tổng kết)\b", re.IGNORECASE),
]


def _is_heading(line: str) -> bool:
    """Nhận diện dòng tiêu đề chương/mục (ngắn, có cấu trúc tiêu đề)."""
    line = (line or "").strip()
    if not line or len(line) > 80:
        return False
    for pattern in _HEADING_PATTERNS:
        if pattern.search(line):
            return True
    # Dòng toàn chữ hoa ngắn (dạng tiêu đề)
    if len(line) <= 60 and re.fullmatch(r"[A-ZÀ-ỸẠ-Ỵ0-9\s&\-:·]+", line):
        return True
    return False


def extract_chapters(text: str) -> list:
    """Chia văn bản thành các chương/mục. Trả về [{title, body: [dòng]}]."""
    lines = [ln.strip() for ln in (text or "").split("\n") if ln.strip()]
    chapters = []
    current = None
    for line in lines:
        if _is_heading(line):
            chapters.append({"title": line, "body": []})
            current = chapters[-1]
        else:
            if current is None:
                if not chapters:
                    chapters.append({"title": "Nội dung chính", "body": []})
                    current = chapters[0]
            if current is not None:
                current["body"].append(line)
    if not chapters:
        chapters.append({"title": "Nội dung chính", "body": []})
    return chapters


# =========================================================================
# SINH FLASHCARD THÔNG MINH THEO CHƯƠNG
# =========================================================================
_DEFINITION_RE = re.compile(
    r"^(.+?)\s+(?:được\s+(?:định nghĩa|hiểu|gọi|coi)\s+là|là|nghĩa là|chính là)\s+(.+)",
    re.IGNORECASE,
)
_SUBJECT_PREFIX_RE = re.compile(
    r"^(theo\s+[^\s,]{1,20}\s*,?\s*|nói\s+chung\s*,?\s*|về\s+[^\s,]{1,20}\s*,?\s*)",
    re.IGNORECASE,
)
# Từ mở đầu câu KHÔNG nên là chủ ngữ của một định nghĩa
_JOINERS = (
    "và", "nhưng", "hoặc", "vì", "để", "theo", "với", "khi", "nếu", "thì", "cũng",
    "đây", "đó", "do", "trong", "ngoài", "trên", "dưới", "từ", "ở", "tại", "qua",
    "về", "cùng", "cho", "đối", "giữa", "sau", "trước", "bằng", "nhờ", "ngoài ra",
)


def _definition_card(sentence: str) -> dict | None:
    """Trích câu định nghĩa kiểu 'X là ...' -> thẻ Q: 'X là gì?' A: nguyên câu."""
    m = _DEFINITION_RE.search(sentence)
    if not m:
        return None
    subject = m.group(1).strip()
    subject = _SUBJECT_PREFIX_RE.sub("", subject).strip()
    subject = subject.rstrip(",:;–—-")
    if len(subject) < 3 or len(subject) > 80:
        return None
    if subject.lower().split()[0] in _JOINERS:
        return None
    return {"question": f"{subject} là gì?", "answer": sentence}


def _blank_word(sentence: str, word: str) -> str:
    """Điền khuyết từ khóa trong câu (khớp nguyên từ, không khớp chuỗi con)."""
    pattern = r"(?<![A-Za-zÀ-ỹ0-9])" + re.escape(word) + r"(?![A-Za-zÀ-ỹ0-9])"
    return re.sub(pattern, "......", sentence, count=1, flags=re.IGNORECASE)


def generate_flashcards(text: str, max_cards: int = 12) -> list:
    """Tự sinh flashcard thông minh theo chương/mục.

    Mỗi chương ưu tiên:
      1. Thẻ ĐỊNH NGHĨA: câu có 'X là ...' -> Q: 'X là gì?'  A: nguyên câu.
      2. Thẻ ĐIỀN KHUYẾT: câu chứa từ khóa quan trọng của chương.

    Trả về [{question, answer, chapter, type, context}].
    """
    raw = (text or "")[:45000]
    if not raw.strip():
        return []

    # Tách chương từ văn bản GỐC (giữ dòng tiêu đề chương/mục)
    chapters = extract_chapters(raw)
    active_chapters = [ch for ch in chapters if len(" ".join(ch["body"]).strip()) >= 30]
    if not active_chapters:
        active_chapters = chapters

    per_chapter = max(1, max_cards // len(active_chapters))
    cards = []
    seen = set()

    for ch in active_chapters:
        body = _clean_text(" ".join(ch["body"]))
        if len(body.strip()) < 30:
            continue
        title = ch["title"]
        top_words = _word_frequencies(body, limit=12)
        sentences = _split_sentences(body)
        ch_cards = []

        # 1) Thẻ định nghĩa
        for s in sentences:
            if len(ch_cards) >= per_chapter:
                break
            if s in seen:
                continue
            dc = _definition_card(s)
            if dc:
                dc.update({"chapter": title, "type": "definition", "context": s})
                ch_cards.append(dc)
                seen.add(s)

        # 2) Thẻ điền khuyết
        for w in top_words:
            if len(ch_cards) >= per_chapter:
                break
            for s in sentences:
                if s in seen:
                    continue
                if w["key"] in _normalize_keyword(s):
                    question = _blank_word(s, w["word"])
                    if question != s:
                        ch_cards.append({
                            "question": question,
                            "answer": w["word"].upper(),
                            "chapter": title,
                            "type": "fill-blank",
                            "context": s,
                        })
                        seen.add(s)
                        break

        cards.extend(ch_cards)
        if len(cards) >= max_cards:
            break

    return cards[:max_cards]


# =========================================================================
# TÍCH HỢP LLM LOCAL (Ollama) — nâng cấp tóm tắt/trích xuất học thuật
# =========================================================================
_LLM_SYSTEM_PROMPT = (
    "Bạn là trợ lý học thuật tiếng Việt cấp đại học, chuyên tóm tắt tài liệu lý thuyết. "
    "Hãy phân tích kỹ nội dung và trả về MỘT JSON hợp lệ duy nhất, không kèm giải thích, không dùng code fence. "
    "Cấu trúc JSON bắt buộc: "
    '{"overview": "tóm tắt tổng quan 3-5 câu: nêu chủ đề, phạm vi và mục tiêu của tài liệu", '
    '"keywords": ["từ khóa chính"], '
    '"chapters": [{"title": "Tên chương/mục", "summary": "tóm tắt nội dung chương 2-4 câu", "keywords": ["từ khóa chương"]}], '
    '"key_points": ["các ý chính quan trọng nhất (5-8 ý)"], '
    '"conclusion": "kết luận hoặc điểm mấu chốt rút ra từ tài liệu", '
    '"terms": [{"term": "thuật ngữ quan trọng", "definition": "định nghĩa ngắn gọn"}], '
    '"flashcards": [{"question": "câu hỏi ôn tập", "answer": "câu trả lời ngắn gọn", "chapter": "Tên chương"}], '
    'sinh 6-10 flashcard ôn tập lý thuyết theo chương}. '
    "Bám sát nội dung tài liệu, không bịa thêm thông tin, ưu tiên các khái niệm và kết luận chính."
)


def _get_llm():
    try:
        from app.services.llm_client import _get_llm as _factory
        return _factory()
    except Exception:
        return None


def analyze_document_ai(raw_text: str, max_chars: int = 9000) -> dict | None:
    """Dùng LLM local để phân tích học thuật. Trả None nếu model không khả dụng/lỗi."""
    llm = _get_llm()
    if llm is None or not llm.available():
        return None
    snippet = _clean_text(raw_text)[:max_chars]
    if not snippet:
        return None
    try:
        from app.services.llm_client import extract_json
        resp = llm.chat(_LLM_SYSTEM_PROMPT, f"Hãy phân tích tài liệu sau:\n\n{snippet}")
        data = extract_json(resp)
        if not data:
            return None

        def _norm_list(v):
            if isinstance(v, list):
                return [str(x).strip() for x in v if str(x).strip()]
            return []

        chapters = []
        for ch in (data.get("chapters") or []):
            if isinstance(ch, dict) and ch.get("title"):
                chapters.append({
                    "title": str(ch["title"]).strip(),
                    "summary": str(ch.get("summary") or "").strip(),
                    "keywords": _norm_list(ch.get("keywords")),
                })

        terms = []
        for t in (data.get("terms") or []):
            if isinstance(t, dict) and t.get("term"):
                terms.append({
                    "term": str(t["term"]).strip(),
                    "definition": str(t.get("definition") or "").strip(),
                })

        return {
            "overview": str(data.get("overview") or "").strip(),
            "keywords": _norm_list(data.get("keywords")),
            "key_points": _norm_list(data.get("key_points")),
            "chapters": chapters,
            "conclusion": str(data.get("conclusion") or "").strip(),
            "terms": terms,
            "flashcards": [
                {
                    "question": str(c.get("question") or "").strip(),
                    "answer": str(c.get("answer") or "").strip(),
                    "chapter": str(c.get("chapter") or "").strip(),
                    "type": "ai",
                }
                for c in (data.get("flashcards") or [])
                if isinstance(c, dict) and c.get("question") and c.get("answer")
            ],
        }
    except Exception:
        return None


def document_analysis(raw_text: str, max_text: int = 45000, use_llm: bool = True) -> dict:
    """Phân tích tài liệu: ưu tiên LLM, fallback về thuật toán quy tắc offline."""
    base = analyze_document(raw_text, max_text)
    if not use_llm:
        return base
    ai = analyze_document_ai(raw_text)
    if not ai:
        return base
    # Gộp kết quả LLM, giữ fallback từng trường nếu thiếu
    if ai.get("overview"):
        base["summary"] = ai["overview"]
    if ai.get("key_points"):
        base["key_points"] = "\n".join(ai["key_points"])
    if ai.get("keywords"):
        kw = ai["keywords"]
        base["keywords"] = ", ".join(kw)
        base["tags"] = ", ".join(kw[:5])
    if ai.get("chapters"):
        base["chapters"] = ai["chapters"]
    if ai.get("conclusion"):
        base["conclusion"] = ai["conclusion"]
    if ai.get("terms"):
        base["terms"] = ai["terms"]
    if ai.get("flashcards"):
        base["_llm_flashcards"] = ai["flashcards"]
    return base


def document_flashcards(raw_text: str, max_cards: int = 12, use_llm: bool = True) -> list:
    """Sinh flashcard: ưu tiên LLM, fallback quy tắc."""
    if use_llm:
        ai = analyze_document_ai(raw_text)
        if ai and ai.get("flashcards"):
            cards = []
            for c in ai["flashcards"][:max_cards]:
                cards.append({
                    "question": c["question"],
                    "answer": c["answer"],
                    "chapter": c.get("chapter") or "",
                    "type": "ai",
                    "context": "",
                })
            return cards
    return generate_flashcards(raw_text, max_cards)


# =========================================================================
# KIỂM DUYỆT TỰ ĐỘNG — toàn bộ vi phạm tài liệu
# =========================================================================
_MODERATION_SYSTEM_PROMPT = (
    "Bạn là bộ kiểm duyệt nội dung tài liệu học thuật. Đánh giá tài liệu theo các loại vi phạm sau:\n"
    '1. "phapluat": vi phạm pháp luật/chính trị (tuyên truyền lật đổ, kích động bạo lực, phỉ báng, nội dung bị cấm)\n'
    '2. "dothi": ngôn từ độc hại (chửi thề, xúc phạm, kỳ thị, đe dọa)\n'
    '3. "doitruy": nội dung đồi trụy/18+ (khiêu dâm, nhạy cảm)\n'
    '4. "spam": quảng cáo, mời tham gia, tuyển dụng, link giới thiệu, bán hàng\n'
    '5. "rip": lộ thông tin cá nhân (CMND, SĐT, email, tài khoản người khác)\n'
    '6. "link_doc_hai": liên kết hoặc file độc hại, phần mềm lạ\n'
    '7. "buon_ban": rao bán luận văn/bài tập, kinh doanh tài liệu trái phép\n'
    '8. "daovan": sao chép nguyên văn không ghi nguồn, vi phạm bản quyền sách\n'
    '9. "sai_chu_de": không liên quan học thuật/môn học (truyện, tin tức, giải trí)\n'
    '10. "chat_luong_thap": scan mờ không đọc được, file rỗng/hỏng, không có nội dung\n'
    '11. "sai_lech": thông tin sai lệch, phản khoa học nghiêm trọng\n'
    'Chỉ trả về một JSON hợp lệ duy nhất, không giải thích: '
    '{"verdict": "approved" | "rejected" | "review", "risk": 0.0-1.0, '
    '"reason": "lý do ngắn gọn tiếng Việt", "categories": ["mã vi phạm"]}. '
    '- "rejected": khi có vi phạm nghiêm trọng (loại 1-7).\n'
    '- "review": khi nghi ngờ vi phạm trung bình/nhẹ (loại 8-11).\n'
    '- "approved": nội dung học thuật bình thường, không vi phạm.'
)


def _moderation_categories_config():
    """Đọc danh sách reject/review categories từ config.yaml."""
    try:
        from config.settings import settings
        mod_cfg = settings.config.get("moderation", {}) or {}
    except Exception:
        mod_cfg = {}
    reject = [str(c).strip() for c in (mod_cfg.get("reject_categories") or []) if str(c).strip()]
    review = [str(c).strip() for c in (mod_cfg.get("review_categories") or []) if str(c).strip()]
    if not reject:
        reject = ["phapluat", "dothi", "doitruy", "spam", "rip", "link_doc_hai", "buon_ban"]
    if not review:
        review = ["daovan", "sai_chu_de", "chat_luong_thap", "sai_lech", "trung_lap"]
    return reject, review


def moderate_content(raw_text: str, max_chars: int = 9000) -> dict:
    """Kiểm duyệt tự động nội dung tài liệu.

    Trả về {verdict, risk, reason, categories}.
    verdict = None khi không kiểm được (LLM không khả dụng) → hệ thống giữ nguyên luồng cũ.
    """
    llm = _get_llm()
    if llm is None or not llm.available():
        return {"verdict": None, "risk": 0.0, "reason": "", "categories": []}
    snippet = _clean_text(raw_text)[:max_chars]
    if not snippet:
        return {"verdict": "approved", "risk": 0.0, "reason": "", "categories": []}
    try:
        from app.services.llm_client import extract_json
        resp = llm.chat(_MODERATION_SYSTEM_PROMPT, f"Kiểm duyệt tài liệu sau:\n\n{snippet}")
        data = extract_json(resp)
        if not data:
            return {"verdict": None, "risk": 0.0, "reason": "", "categories": []}
        verdict = str(data.get("verdict") or "approved").strip().lower()
        if verdict not in ("approved", "rejected", "review"):
            verdict = "approved"
        categories = [str(c).strip() for c in (data.get("categories") or []) if str(c).strip()]
        risk = float(data.get("risk") or 0.0)
        reason = str(data.get("reason") or "").strip()

        # Áp dụng quy tắc chặt: nếu chứa mã reject/review thì ép verdict tương ứng
        reject, review = _moderation_categories_config()
        if any(c in reject for c in categories):
            verdict = "rejected"
            if risk < 0.7:
                risk = 0.9
        elif any(c in review for c in categories):
            verdict = "review"
            if risk < 0.4:
                risk = 0.5

        return {
            "verdict": verdict,
            "risk": risk,
            "reason": reason,
            "categories": categories,
        }
    except Exception:
        return {"verdict": None, "risk": 0.0, "reason": "", "categories": []}
